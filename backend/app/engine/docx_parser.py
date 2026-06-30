import docx
import logging

logger = logging.getLogger("ats-optimizer")

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(file_path)
        text_runs = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_runs.append(paragraph.text)
                
        # Extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_runs.append(cell.text)
                        
        return "\n".join(text_runs)
    except Exception as e:
        logger.error(f"python-docx extraction failed: {e}")
        raise Exception("Failed to extract text from DOCX file.")
