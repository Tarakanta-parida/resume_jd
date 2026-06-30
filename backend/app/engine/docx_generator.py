import io
import re
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def clean_html_tags(text: str) -> str:
    if not text:
        return ""
    return re.sub(r'</?[a-zA-Z]+(?:\s+[^>]*)?>', '', text)


def generate_docx_from_data(data: dict) -> io.BytesIO:
    doc = docx.Document()
    
    # Set page margins (0.75 inch is standard and fits everything on one page better)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Helper for heading styling
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        
    # Helper for body paragraph
    def add_body_paragraph(text, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        return p

    # 1. Personal Info (Header)
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    run_name = p_header.add_run(data["personalInfo"]["name"].upper())
    run_name.bold = True
    run_name.font.size = Pt(16)
    run_name.font.name = 'Arial'
    
    # Sub-header (contact info)
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    contact_parts = []
    if data["personalInfo"].get("email"):
        contact_parts.append(data["personalInfo"]["email"])
    if data["personalInfo"].get("phone"):
        contact_parts.append(data["personalInfo"]["phone"])
    if data["personalInfo"].get("linkedin"):
        contact_parts.append(data["personalInfo"]["linkedin"])
    if data["personalInfo"].get("github"):
        contact_parts.append(data["personalInfo"]["github"])
    
    run_contact = p_contact.add_run("  |  ".join(contact_parts))
    run_contact.font.size = Pt(9.5)
    run_contact.font.name = 'Arial'
    
    # 2. Professional Summary
    if data.get("summary"):
        add_section_heading("Professional Summary")
        clean_summary = clean_html_tags(data["summary"])
        add_body_paragraph(clean_summary, space_after=12)
        
    # 3. Technical Skills (Skills)
    if data.get("skills"):
        add_section_heading("Technical Skills")
        clean_skills = []
        for s in data["skills"]:
            s_clean = clean_html_tags(s)
            clean_skills.append(s_clean)
        skills_text = "  •  ".join(clean_skills)
        add_body_paragraph(skills_text, space_after=12)
        
    # 4. Experience
    if data.get("experience"):
        add_section_heading("Work Experience")
        for exp in data["experience"]:
            # Header line: Role | Company (Duration)
            role_clean = clean_html_tags(exp["role"])
            
            p_exp_header = doc.add_paragraph()
            p_exp_header.paragraph_format.space_after = Pt(1)
            p_exp_header.paragraph_format.keep_with_next = True
            
            run_role = p_exp_header.add_run(f"{role_clean}  |  ")
            run_role.bold = True
            run_role.font.name = 'Arial'
            run_role.font.size = Pt(11)
            
            run_company = p_exp_header.add_run(f"{exp['company']}")
            run_company.font.italic = True
            run_company.font.name = 'Arial'
            run_company.font.size = Pt(10.5)
            
            run_duration = p_exp_header.add_run(f" ({exp['duration']})")
            run_duration.font.name = 'Arial'
            run_duration.font.size = Pt(10)
            
            # Bullets
            for bullet in exp.get("bullets", []):
                bullet_clean = clean_html_tags(bullet)
                
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.paragraph_format.line_spacing = 1.15
                run_bullet = p_bullet.add_run(bullet_clean)
                run_bullet.font.name = 'Arial'
                run_bullet.font.size = Pt(10.5)
                
    # 5. Technical Projects
    if data.get("projects") and len(data["projects"]) > 0:
        non_empty_projects = [p for p in data["projects"] if p.get("name")]
        if len(non_empty_projects) > 0:
            add_section_heading("Technical Projects")
            for proj in non_empty_projects:
                desc_clean = clean_html_tags(proj["description"])
                
                p_proj_header = doc.add_paragraph()
                p_proj_header.paragraph_format.space_after = Pt(1)
                p_proj_header.paragraph_format.keep_with_next = True
                
                run_proj_name = p_proj_header.add_run(proj["name"])
                run_proj_name.bold = True
                run_proj_name.font.name = 'Arial'
                run_proj_name.font.size = Pt(11)
                
                if desc_clean:
                    run_proj_desc = p_proj_header.add_run(f"  |  {desc_clean}")
                    run_proj_desc.font.italic = True
                    run_proj_desc.font.name = 'Arial'
                    run_proj_desc.font.size = Pt(10.5)
                    
                for bullet in proj.get("bullets", []):
                    bullet_clean = clean_html_tags(bullet)
                    
                    p_bullet = doc.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_after = Pt(2)
                    p_bullet.paragraph_format.line_spacing = 1.15
                    run_bullet = p_bullet.add_run(bullet_clean)
                    run_bullet.font.name = 'Arial'
                    run_bullet.font.size = Pt(10.5)

    # 6. Education
    if data.get("education"):
        add_section_heading("Education")
        for edu in data["education"]:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = Pt(3)
            
            run_degree = p_edu.add_run(f"{edu['degree']}  |  ")
            run_degree.bold = True
            run_degree.font.name = 'Arial'
            run_degree.font.size = Pt(11)
            
            run_school = p_edu.add_run(f"{edu['school']}")
            run_school.font.italic = True
            run_school.font.name = 'Arial'
            run_school.font.size = Pt(10.5)
            
            run_year = p_edu.add_run(f" ({edu['year']})")
            run_year.font.name = 'Arial'
            run_year.font.size = Pt(10)

    # 7. Certifications
    if data.get("certifications"):
        add_section_heading("Certifications")
        for cert in data["certifications"]:
            cert_clean = clean_html_tags(cert)
            p_cert = doc.add_paragraph(style='List Bullet')
            p_cert.paragraph_format.space_after = Pt(2)
            p_cert.paragraph_format.line_spacing = 1.15
            run_cert = p_cert.add_run(cert_clean)
            run_cert.font.name = 'Arial'
            run_cert.font.size = Pt(10.5)

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
